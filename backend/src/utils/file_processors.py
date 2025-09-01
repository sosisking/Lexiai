import os
import io
import tempfile
import subprocess
from flask import current_app

class FileProcessor:
    """Base class for file processors."""
    
    @staticmethod
    def get_processor(file_type):
        """
        Get the appropriate processor for a file type.
        
        Args:
            file_type (str): The file type
            
        Returns:
            FileProcessor: The processor
        """
        processors = {
            'pdf': PDFProcessor,
            'docx': DocxProcessor,
            'txt': TextProcessor
        }
        
        return processors.get(file_type.lower(), TextProcessor)
    
    @staticmethod
    def extract_text(file_path):
        """
        Extract text from a file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            str: The extracted text
        """
        raise NotImplementedError("Subclasses must implement extract_text")
    
    @staticmethod
    def extract_metadata(file_path):
        """
        Extract metadata from a file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            dict: The metadata
        """
        raise NotImplementedError("Subclasses must implement extract_metadata")


class PDFProcessor(FileProcessor):
    """Processor for PDF files."""
    
    @staticmethod
    def extract_text(file_path):
        """
        Extract text from a PDF file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            str: The extracted text
        """
        try:
            # Use pdftotext (from poppler-utils) to extract text
            result = subprocess.run(
                ['pdftotext', file_path, '-'],
                capture_output=True,
                text=True,
                check=True
            )
            return result.stdout
        except subprocess.CalledProcessError as e:
            current_app.logger.error(f"Error extracting text from PDF: {str(e)}")
            return ""
    
    @staticmethod
    def extract_metadata(file_path):
        """
        Extract metadata from a PDF file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            dict: The metadata
        """
        try:
            # Use pdfinfo (from poppler-utils) to extract metadata
            result = subprocess.run(
                ['pdfinfo', file_path],
                capture_output=True,
                text=True,
                check=True
            )
            
            # Parse metadata
            metadata = {}
            for line in result.stdout.splitlines():
                if ':' in line:
                    key, value = line.split(':', 1)
                    metadata[key.strip()] = value.strip()
            
            return metadata
        except subprocess.CalledProcessError as e:
            current_app.logger.error(f"Error extracting metadata from PDF: {str(e)}")
            return {}


class DocxProcessor(FileProcessor):
    """Processor for DOCX files."""
    
    @staticmethod
    def extract_text(file_path):
        """
        Extract text from a DOCX file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            str: The extracted text
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            return '\n'.join(text)
        except Exception as e:
            current_app.logger.error(f"Error extracting text from DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_metadata(file_path):
        """
        Extract metadata from a DOCX file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            dict: The metadata
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            metadata = {}
            
            # Extract core properties
            core_props = doc.core_properties
            metadata['title'] = core_props.title
            metadata['author'] = core_props.author
            metadata['created'] = core_props.created.isoformat() if core_props.created else None
            metadata['modified'] = core_props.modified.isoformat() if core_props.modified else None
            
            return metadata
        except Exception as e:
            current_app.logger.error(f"Error extracting metadata from DOCX: {str(e)}")
            return {}


class TextProcessor(FileProcessor):
    """Processor for text files."""
    
    @staticmethod
    def extract_text(file_path):
        """
        Extract text from a text file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            str: The extracted text
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            current_app.logger.error(f"Error extracting text from text file: {str(e)}")
            return ""
    
    @staticmethod
    def extract_metadata(file_path):
        """
        Extract metadata from a text file.
        
        Args:
            file_path (str): The file path
            
        Returns:
            dict: The metadata
        """
        try:
            stat = os.stat(file_path)
            
            metadata = {
                'size': stat.st_size,
                'created': stat.st_ctime,
                'modified': stat.st_mtime
            }
            
            return metadata
        except Exception as e:
            current_app.logger.error(f"Error extracting metadata from text file: {str(e)}")
            return {}

